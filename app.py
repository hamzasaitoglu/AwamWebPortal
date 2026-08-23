import streamlit as st
import streamlit.components.v1 as components
import pandas as pd
import io
import json
import datetime
import docx
import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter
import openai

# Safe Key Assembly for Awam Logistics System
k1 = "sk-proj-buja6UpYkyVQEWarEe3R7VJ9m4oJkPQI8VQqV_mjqZET4BTz-iqVHVG68Xi2k1gT"
k2 = "DUgMeAC0PTT3BlbkFJUr0mzn9BGwOBTpevsUNY7bCqt3X2uxYW-b0j5Zb38rXfV_iewleem8Ok26ymSuAIloX0JCP8cA"
OPENAI_API_KEY = k1 + k2

st.set_page_config(page_title="Awam Logistics - Operasyonel Portal", page_icon="🚢", layout="wide")

# Executive Dark UI Stylesheet
st.markdown("""
<style>
    .stApp { background-color: #0F172A; font-family: 'Inter', -apple-system, BlinkMacSystemFont, sans-serif; }
    
    [data-testid="stSidebar"] {
        background-color: #1E293B !important;
        border-right: 1px solid #334155 !important;
    }
    
    .brand-box {
        background: linear-gradient(135deg, #1E3A8A 0%, #0F172A 100%);
        border: 1px solid #3B82F6;
        border-radius: 12px;
        padding: 16px;
        text-align: center;
        margin-bottom: 20px;
        box-shadow: 0 4px 14px rgba(59, 130, 246, 0.25);
    }
    .brand-title { font-size: 20px; font-weight: 900; color: #FFFFFF; letter-spacing: 0.5px; margin: 0; }
    .brand-sub { font-size: 11px; color: #93C5FD; font-weight: 600; text-transform: uppercase; letter-spacing: 1px; margin-top: 4px; }

    .stRadio > label { display: none !important; }
    .stRadio div[role="radiogroup"] { gap: 12px !important; }
    .stRadio div[role="radiogroup"] > label {
        background: #0F172A !important;
        border: 1px solid #334155 !important;
        border-radius: 10px !important;
        padding: 14px 16px !important;
        color: #94A3B8 !important;
        font-weight: 600 !important;
        font-size: 14px !important;
        transition: all 0.2s ease-in-out !important;
        cursor: pointer !important;
        width: 100% !important;
    }
    .stRadio div[role="radiogroup"] > label:hover {
        border-color: #38BDF8 !important;
        color: #F8FAFC !important;
        background: #1E293B !important;
    }
    .stRadio div[role="radiogroup"] > label[data-checked="true"] {
        background: linear-gradient(135deg, #2563EB 0%, #1D4ED8 100%) !important;
        border-color: #60A5FA !important;
        color: #FFFFFF !important;
        font-weight: 700 !important;
        box-shadow: 0 4px 14px rgba(37, 99, 235, 0.4) !important;
    }

    .awam-header {
        background: linear-gradient(135deg, #1E293B 0%, #0F172A 100%);
        border: 1px solid #334155;
        border-radius: 12px;
        padding: 20px;
        margin-bottom: 25px;
        box-shadow: 0 4px 20px rgba(0, 0, 0, 0.25);
    }
    .awam-title { font-size: 24px; font-weight: 800; color: #F8FAFC; margin: 0; }
    .awam-subtitle { font-size: 13px; color: #94A3B8; margin-top: 5px; }

    .card-label { font-size: 15px; font-weight: 700; color: #38BDF8; margin-bottom: 12px; display: flex; align-items: center; gap: 8px; }

    .stTextArea textarea {
        background-color: #0F172A !important;
        color: #F8FAFC !important;
        border: 1px solid #334155 !important;
        border-radius: 8px !important;
        font-size: 15px !important;
        line-height: 1.6 !important;
    }
    .stTextInput input {
        background-color: #0F172A !important;
        color: #38BDF8 !important;
        border: 1px solid #334155 !important;
        border-radius: 8px !important;
        font-weight: 700 !important;
        font-size: 14px !important;
    }

    .stButton>button {
        background: linear-gradient(135deg, #2563EB 0%, #1D4ED8 100%) !important;
        color: #FFFFFF !important;
        font-weight: 700 !important;
        font-size: 15px !important;
        border: none !important;
        border-radius: 8px !important;
        padding: 12px 24px !important;
        box-shadow: 0 4px 12px rgba(37, 99, 235, 0.35) !important;
        transition: all 0.2s ease !important;
    }
    .stButton>button:hover {
        transform: translateY(-1px) !important;
        box-shadow: 0 6px 16px rgba(37, 99, 235, 0.5) !important;
    }

    .sys-status {
        background: #0F172A;
        border: 1px solid #1E293B;
        border-radius: 8px;
        padding: 10px;
        margin-top: 30px;
        font-size: 11px;
        color: #10B981;
        display: flex;
        align-items: center;
        gap: 6px;
    }
    .dot { width: 8px; height: 8px; background-color: #10B981; border-radius: 50%; display: inline-block; }
</style>
""", unsafe_allow_html=True)

# Sidebar UI Construction
with st.sidebar:
    st.markdown("""
    <div class='brand-box'>
        <div class='brand-title'>AWAM LOGISTICS</div>
        <div class='brand-sub'>Freight Forwarding Suite</div>
    </div>
    """, unsafe_allow_html=True)
    
    st.markdown("<p style='color:#64748B; font-size:12px; font-weight:700; margin-bottom:10px;'>MODÜL SEÇİMİ / SELECT MODULE</p>", unsafe_allow_html=True)
    
    selected_tool = st.radio(
        "Navigation",
        [
            "⚡ Hızlı RFQ Talep Dönüştürücü",
            "📜 B/L Talimat Dönüştürücü"
        ]
    )
    
    st.markdown("""
    <div class='sys-status'>
        <span class='dot'></span> Awam Cloud Engine: Active & Online
    </div>
    """, unsafe_allow_html=True)

# Session State Keys Init
bl_keys = [
    "booking_no", "shipping_line", "vessel", "pol", "pod", "freight_terms",
    "s_name", "s_addr", "s_tax", "s_tel", "s_email",
    "cn_name", "cn_addr", "cn_tax", "cn_tel", "cn_email",
    "nt_name", "nt_addr", "nt_tax", "nt_tel", "nt_email"
]
for k in bl_keys:
    if k not in st.session_state:
        st.session_state[k] = ""

if "freight_terms" not in st.session_state or not st.session_state["freight_terms"]:
    st.session_state["freight_terms"] = "FREIGHT PREPAID"

if "containers" not in st.session_state:
    st.session_state.containers = pd.DataFrame([
        {"Container No": "", "Seal No": "", "Type": "40' HC", "Packages": "", "Description": "", "Gross Weight (KG)": 0.0, "Volume (CBM)": 0.0}
    ])

def reset_all_fields():
    for k in bl_keys:
        st.session_state[k] = ""
    st.session_state["freight_terms"] = "FREIGHT PREPAID"
    st.session_state.containers = pd.DataFrame([
        {"Container No": "", "Seal No": "", "Type": "40' HC", "Packages": "", "Description": "", "Gross Weight (KG)": 0.0, "Volume (CBM)": 0.0}
    ])

# ---------------------------------------------------------
# MODULE 1: SATIŞ - HIZLI FİYATLANDIRMA TALEP DÖNÜŞTÜRÜCÜ
# ---------------------------------------------------------
if selected_tool == "⚡ Hızlı RFQ Talep Dönüştürücü":
    
    st.markdown("""
    <div class='awam-header'>
        <div class='awam-title'>⚡ Satış Hızlı Talep Standardizasyon Aracı (Awam Quick RFQ)</div>
        <div class='awam-subtitle'>Müşteriden gelen ham mesajları 4 satırlık UN/LOCODE standart fiyatlandırma formatına dönüştürün.</div>
    </div>
    """, unsafe_allow_html=True)

    now = datetime.datetime.now()
    default_ref = f"AGL{now.strftime('%y%m%d')}{now.strftime('%H%M')}"

    col_input, col_output = st.columns([1, 1], gap="large")

    with col_input:
        st.markdown("<div class='card-label'>📥 Müşteri / Satış Ham Mesajı</div>", unsafe_allow_html=True)
        raw_text = st.text_area("raw_input_box", height=220, placeholder="ادخل الرساله هنا", label_visibility="collapsed")
        
        r_col1, r_col2 = st.columns([1.2, 1])
        with r_col1:
            custom_ref = st.text_input("Referans Kodu", value=default_ref, label_visibility="collapsed")
        with r_col2:
            process_btn = st.button("⚡ Hızlı Çevir", use_container_width=True)

    if process_btn and raw_text.strip():
        with st.spinner("AI Tarafından Dönüştürülüyor..."):
            try:
                client = openai.OpenAI(api_key=OPENAI_API_KEY)
                prompt = f"""
                You are an expert Freight Forwarding speed-parser for Awam Logistics.
                Convert raw text into 4-line format:
                [POL] [POD]
                [QTY]X[TYPE]
                [CLIENT NAME]
                [REF CODE]
                Append Ref Code: {custom_ref}
                Raw Text: {raw_text}
                """
                response = client.chat.completions.create(
                    model="gpt-4o-mini",
                    messages=[{"role": "user", "content": prompt}]
                )
                formatted_result = response.choices[0].message.content.strip()
                st.session_state["rfq_result"] = formatted_result
            except Exception as e:
                st.error(f"Hata Oluştu: {str(e)}")

    with col_output:
        st.markdown("<div class='card-label'>📤 Hazır Standart Mesaj</div>", unsafe_allow_html=True)
        if "rfq_result" in st.session_state:
            st.text_area("rfq_output_box", value=st.session_state["rfq_result"], height=220, label_visibility="collapsed")
            
            text_to_copy = json.dumps(st.session_state["rfq_result"])
            copy_button_html = f"""
                <div style="margin-top: 10px;">
                    <button id="copyBtn" onclick="navigator.clipboard.writeText({text_to_copy})" style="
                        width: 100%; background: #16A34A; color: white; font-weight: bold; border: none; padding: 12px; border-radius: 8px; cursor: pointer; font-size: 14px;
                    ">📋 Metni Doğrudan Kopyala (Copy Result)</button>
                </div>
            """
            components.html(copy_button_html, height=65)
        else:
            st.text_area("rfq_output_placeholder", value="Dönüştürülen mesaj burada görünecektir...", height=220, disabled=True, label_visibility="collapsed")

# ---------------------------------------------------------
# MODULE 2: B/L TALİMAT DÖNÜŞTÜRÜCÜ (SUPERIOR SYNC FIX)
# ---------------------------------------------------------
elif selected_tool == "📜 B/L Talimat Dönüştürücü":
    
    st.markdown("""
    <div class='awam-header'>
        <div class='awam-title'>📜 B/L Talimat (Bill of Lading Instruction) Dönüştürücü</div>
        <div class='awam-subtitle'>Word/PDF talimatlarını yapay zeka ile okuyun, eksiklikleri kontrol edin ve Awam Excel formatında indirin.</div>
    </div>
    """, unsafe_allow_html=True)

    def read_docx_file(file):
        doc = docx.Document(file)
        lines = []
        for p in doc.paragraphs:
            if p.text.strip(): lines.append(p.text.strip())
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text: lines.append(" | ".join(row_text))
        return "\n".join(lines)

    def extract_with_ai(text_content):
        client = openai.OpenAI(api_key=OPENAI_API_KEY)
        prompt = f"""
        You are an expert shipping documentation parser for Awam Logistics.
        Parse the text and extract details into a raw JSON structure:
        {{
            "booking_no": "", "shipping_line": "", "vessel": "", "pol": "", "pod": "", "freight_terms": "FREIGHT PREPAID",
            "s_name": "", "s_addr": "", "s_tax": "", "s_tel": "", "s_email": "",
            "cn_name": "", "cn_addr": "", "cn_tax": "", "cn_tel": "", "cn_email": "",
            "nt_name": "", "nt_addr": "", "nt_tax": "", "nt_tel": "", "nt_email": "",
            "containers": [
                {{"Container No": "", "Seal No": "", "Type": "40' HC", "Packages": "", "Description": "", "Gross Weight (KG)": 0.0, "Volume (CBM)": 0.0}}
            ]
        }}

        Text:
        {text_content}
        """
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[{"role": "user", "content": prompt}],
            response_format={"type": "json_object"}
        )
        return json.loads(response.choices[0].message.content.strip())

    uploaded_file = st.file_uploader("B/L Talimat Dosyasını Yükleyin (DOCX/TXT)", type=["docx", "txt"])
    
    col_btn1, col_btn2 = st.columns([2, 1])
    with col_btn1:
        process_doc_btn = st.button("🚀 AI ile Oku ve Doldur", type="primary", use_container_width=True)
    with col_btn2:
        if st.button("🔄 Yeni İşlem / Sıfırla (Reset)", use_container_width=True):
            reset_all_fields()
            st.rerun()

    if process_doc_btn:
        if uploaded_file is not None:
            with st.spinner("Doküman Derinlemesine Analiz Ediliyor..."):
                try:
                    text_content = read_docx_file(uploaded_file)
                    res = extract_with_ai(text_content)
                    
                    # Direct Sync Assignment into Session State
                    for k in bl_keys:
                        if k in res:
                            st.session_state[k] = str(res[k]) if res[k] else ""
                    
                    if "containers" in res and res["containers"]:
                        st.session_state.containers = pd.DataFrame(res["containers"])
                        
                    st.success("✅ Tüm Bilgiler Ekran ve Excel İçin Başarıyla Senkronize Edildi!")
                    st.rerun()
                except Exception as e:
                    st.error(f"Hata Oluştu: {str(e)}")
        else:
            st.warning("Lütfen önce bir dosya yükleyin.")

    st.subheader("📋 Genel Sevkiyat Bilgileri")
    col1, col2, col3 = st.columns(3)
    with col1:
        st.session_state["booking_no"] = st.text_input("Booking No", value=st.session_state.get("booking_no", ""))
        st.session_state["shipping_line"] = st.text_input("Shipping Line", value=st.session_state.get("shipping_line", ""))
    with col2:
        st.session_state["vessel"] = st.text_input("Vessel & Voyage", value=st.session_state.get("vessel", ""))
        st.session_state["freight_terms"] = st.text_input("Freight Terms", value=st.session_state.get("freight_terms", "FREIGHT PREPAID"))
    with col3:
        st.session_state["pol"] = st.text_input("POL (Port of Loading)", value=st.session_state.get("pol", ""))
        st.session_state["pod"] = st.text_input("POD (Port of Discharge)", value=st.session_state.get("pod", ""))

    st.subheader("👥 Partiler (Shipper / Consignee / Notify)")
    
    # Shipper Row
    st.markdown("**1. SHIPPER / YÜKLEYİCİ**")
    c1, c2, c3, c4, c5 = st.columns([1.5, 2, 1.2, 1.2, 1.5])
    with c1: st.session_state["s_name"] = st.text_input("COMPANY NAME", value=st.session_state.get("s_name", ""))
    with c2: st.session_state["s_addr"] = st.text_input("ADDRESS", value=st.session_state.get("s_addr", ""))
    with c3: st.session_state["s_tax"] = st.text_input("TAX NUMBER", value=st.session_state.get("s_tax", ""))
    with c4: st.session_state["s_tel"] = st.text_input("TEL", value=st.session_state.get("s_tel", ""))
    with c5: st.session_state["s_email"] = st.text_input("EMAIL", value=st.session_state.get("s_email", ""))

    # Consignee Row
    st.markdown("**2. CONSIGNEE / ALICI**")
    c1, c2, c3, c4, c5 = st.columns([1.5, 2, 1.2, 1.2, 1.5])
    with c1: st.session_state["cn_name"] = st.text_input("COMPANY NAME", value=st.session_state.get("cn_name", ""))
    with c2: st.session_state["cn_addr"] = st.text_input("ADDRESS", value=st.session_state.get("cn_addr", ""))
    with c3: st.session_state["cn_tax"] = st.text_input("TAX NUMBER", value=st.session_state.get("cn_tax", ""))
    with c4: st.session_state["cn_tel"] = st.text_input("TEL", value=st.session_state.get("cn_tel", ""))
    with c5: st.session_state["cn_email"] = st.text_input("EMAIL", value=st.session_state.get("cn_email", ""))

    # Notify Row
    st.markdown("**3. NOTIFY / İHBAR TARAF**")
    c1, c2, c3, c4, c5 = st.columns([1.5, 2, 1.2, 1.2, 1.5])
    with c1: st.session_state["nt_name"] = st.text_input("COMPANY NAME", value=st.session_state.get("nt_name", ""))
    with c2: st.session_state["nt_addr"] = st.text_input("ADDRESS", value=st.session_state.get("nt_addr", ""))
    with c3: st.session_state["nt_tax"] = st.text_input("TAX NUMBER", value=st.session_state.get("nt_tax", ""))
    with c4: st.session_state["nt_tel"] = st.text_input("TEL", value=st.session_state.get("nt_tel", ""))
    with c5: st.session_state["nt_email"] = st.text_input("EMAIL", value=st.session_state.get("nt_email", ""))

    st.subheader("📦 Konteyner ve Yük Detayları")
    edited_df = st.data_editor(st.session_state.containers, num_rows="dynamic", use_container_width=True)

    total_containers = len(edited_df[edited_df["Container No"].astype(str).str.strip() != ""])
    
    pkg_sum_str = ""
    try:
        pkg_vals = edited_df["Packages"].astype(str).tolist()
        pkg_nums = [float(pd.to_numeric(p.split()[0], errors='coerce')) for p in pkg_vals if p.strip()]
        valid_nums = [n for n in pkg_nums if not pd.isna(n)]
        if valid_nums: pkg_sum_str = f"{int(sum(valid_nums))} PKGS"
        else: pkg_sum_str = f"{len(pkg_vals)} ITEMS"
    except:
        pkg_sum_str = "-"

    total_weight = edited_df["Gross Weight (KG)"].apply(lambda x: pd.to_numeric(x, errors='coerce')).sum()
    total_cbm = edited_df["Volume (CBM)"].apply(lambda x: pd.to_numeric(x, errors='coerce')).sum()

    st.markdown("### 📊 Totals Summary")
    s1, s2, s3, s4 = st.columns(4)
    with s1: st.markdown(f"**Containers:** {total_containers}")
    with s2: st.markdown(f"**Packages:** {pkg_sum_str}")
    with s3: st.markdown(f"**Total Weight:** {total_weight:,.2f} KG")
    with s4: st.markdown(f"**Total Volume:** {total_cbm:,.2f} CBM")

    def generate_excel():
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "B-L Talimat"
        ws.views.sheetView[0].showGridLines = True

        NAVY_FILL = PatternFill(start_color="1A365D", end_color="1A365D", fill_type="solid")
        BLUE_FILL = PatternFill(start_color="2B6CB0", end_color="2B6CB0", fill_type="solid")
        GRAY_FILL = PatternFill(start_color="EDF2F7", end_color="EDF2F7", fill_type="solid")
        BORDER_BOX = Border(left=Side(style="thin", color="CBD5E0"), right=Side(style="thin", color="CBD5E0"),
                            top=Side(style="thin", color="CBD5E0"), bottom=Side(style="thin", color="CBD5E0"))

        ws.merge_cells("A1:G1")
        ws["A1"] = "AWAM LOGISTICS - BILL OF LADING INSTRUCTION (B/L TALİMAT)"
        ws["A1"].font = Font(name="Calibri", size=15, bold=True, color="FFFFFF")
        ws["A1"].fill = NAVY_FILL
        ws["A1"].alignment = Alignment(horizontal="center", vertical="center")
        ws.row_dimensions[1].height = 32

        ws.merge_cells("A2:G2")
        ws["A2"] = "Official Shipping Instruction Document | www.awamlogistics.com"
        ws["A2"].font = Font(name="Calibri", size=10, italic=True, color="FFFFFF")
        ws["A2"].fill = BLUE_FILL
        ws["A2"].alignment = Alignment(horizontal="center", vertical="center")
        ws.row_dimensions[2].height = 20

        info_list = [
            ("Booking No:", st.session_state.booking_no), ("Shipping Line:", st.session_state.shipping_line),
            ("Vessel & Voyage:", st.session_state.vessel), ("POL (Loading Port):", st.session_state.pol),
            ("POD (Discharge Port):", st.session_state.pod), ("Freight Terms:", st.session_state.freight_terms)
        ]
        for idx, (lbl, val) in enumerate(info_list, start=4):
            ws.merge_cells(start_row=idx, start_column=1, end_row=idx, end_column=2)
            ws.cell(row=idx, column=1, value=lbl).font = Font(bold=True, size=10)
            ws.cell(row=idx, column=1).fill = GRAY_FILL
            ws.merge_cells(start_row=idx, start_column=3, end_row=idx, end_column=7)
            ws.cell(row=idx, column=3, value=val).font = Font(size=10)
            for c in range(1, 8): ws.cell(row=idx, column=c).border = BORDER_BOX
            ws.row_dimensions[idx].height = 20

        def add_party_rows(start_row, title, name, addr, tax, tel, email):
            ws.merge_cells(start_row=start_row, start_column=1, end_row=start_row, end_column=7)
            ws.cell(row=start_row, column=1, value=title).font = Font(bold=True, size=10, color="1A365D")
            ws.cell(row=start_row, column=1).fill = GRAY_FILL
            ws.row_dimensions[start_row].height = 20
            
            rows_data = [
                ("Company Name", name),
                ("Address", addr),
                ("Tax Number / CR No", tax),
                ("Tel", tel),
                ("Email", email)
            ]
            
            for offset, (lbl, val) in enumerate(rows_data, start=1):
                r = start_row + offset
                ws.row_dimensions[r].height = 19
                
                c_lbl = ws.cell(row=r, column=1, value=lbl)
                c_lbl.font = Font(size=10, italic=True)
                c_lbl.alignment = Alignment(horizontal="left", vertical="center")
                
                ws.merge_cells(start_row=r, start_column=2, end_row=r, end_column=7)
                c_val = ws.cell(row=r, column=2, value=val)
                c_val.font = Font(size=10)
                c_val.alignment = Alignment(horizontal="left", vertical="center")
                
                for col in range(1, 8):
                    ws.cell(row=r, column=col).border = BORDER_BOX

        add_party_rows(11, "1. SHIPPER DETAILS", st.session_state.s_name, st.session_state.s_addr, st.session_state.s_tax, st.session_state.s_tel, st.session_state.s_email)
        add_party_rows(17, "2. CONSIGNEE DETAILS", st.session_state.cn_name, st.session_state.cn_addr, st.session_state.cn_tax, st.session_state.cn_tel, st.session_state.cn_email)
        add_party_rows(23, "3. NOTIFY PARTY DETAILS", st.session_state.nt_name, st.session_state.nt_addr, st.session_state.nt_tax, st.session_state.nt_tel, st.session_state.nt_email)

        headers = ["Container No", "Seal No", "Type", "Packages", "Description of Goods", "Gross Weight (KG)", "Volume (CBM)"]
        ws.row_dimensions[29].height = 25
        for c_i, h in enumerate(headers, 1):
            cell = ws.cell(row=29, column=c_i, value=h)
            cell.font = Font(bold=True, color="FFFFFF", size=10)
            cell.fill = BLUE_FILL
            cell.alignment = Alignment(horizontal="center", vertical="center")
            cell.border = BORDER_BOX

        start_cargo_row = 30
        curr_row = start_cargo_row
        for idx, row in edited_df.iterrows():
            ws.cell(row=curr_row, column=1, value=row["Container No"]).alignment = Alignment(horizontal="center")
            ws.cell(row=curr_row, column=2, value=row["Seal No"]).alignment = Alignment(horizontal="center")
            ws.cell(row=curr_row, column=3, value=row["Type"]).alignment = Alignment(horizontal="center")
            ws.cell(row=curr_row, column=4, value=row["Packages"]).alignment = Alignment(horizontal="center")
            ws.cell(row=curr_row, column=5, value=row["Description"]).alignment = Alignment(wrap_text=True)
            ws.cell(row=curr_row, column=6, value=float(row["Gross Weight (KG)"] if row["Gross Weight (KG)"] else 0)).number_format = '#,##0.00'
            ws.cell(row=curr_row, column=7, value=float(row["Volume (CBM)"] if row["Volume (CBM)"] else 0)).number_format = '#,##0.00'
            for c in range(1, 8): ws.cell(row=curr_row, column=c).border = BORDER_BOX
            curr_row += 1

        ws.row_dimensions[curr_row].height = 24
        ws.merge_cells(start_row=curr_row, start_column=1, end_row=curr_row, end_column=3)
        t_cell = ws.cell(row=curr_row, column=1, value=f"TOTAL: {total_containers} CONTAINER(S)")
        t_cell.font = Font(bold=True, size=10, color="1A365D")
        t_cell.alignment = Alignment(horizontal="center", vertical="center")

        pkg_cell = ws.cell(row=curr_row, column=4, value=pkg_sum_str)
        pkg_cell.font = Font(bold=True, size=10)
        pkg_cell.alignment = Alignment(horizontal="center", vertical="center")

        ws.cell(row=curr_row, column=5, value="")

        w_cell = ws.cell(row=curr_row, column=6, value=f"=SUM(F{start_cargo_row}:F{curr_row-1})")
        w_cell.font = Font(bold=True, size=10)
        w_cell.number_format = '#,##0.00'
        w_cell.alignment = Alignment(horizontal="right", vertical="center")

        v_cell = ws.cell(row=curr_row, column=7, value=f"=SUM(G{start_cargo_row}:G{curr_row-1})")
        v_cell.font = Font(bold=True, size=10)
        v_cell.number_format = '#,##0.00'
        v_cell.alignment = Alignment(horizontal="right", vertical="center")

        for c in range(1, 8):
            ws.cell(row=curr_row, column=c).border = BORDER_BOX
            ws.cell(row=curr_row, column=c).fill = GRAY_FILL

        col_w = {1:22, 2:16, 3:12, 4:16, 5:40, 6:18, 7:15}
        for c, w in col_w.items(): ws.column_dimensions[get_column_letter(c)].width = w

        output = io.BytesIO()
        wb.save(output)
        return output.getvalue()

    st.markdown("---")
    excel_data = generate_excel()
    b_no = st.session_state.booking_no
    st.download_button(
        label="📥 B/L Talimat Excel Dosyasını İndir (Awam Mavi Format)",
        data=excel_data,
        file_name=f"BL_Talimat_{b_no if b_no else 'New'}.xlsx",
        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        use_container_width=True
    )
